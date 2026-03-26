"""
Export tools for generating PDF, Excel, Word, and CSV files.
"""

from __future__ import annotations

import csv
import json
import logging
import os
import uuid
from typing import Any

from langchain_core.tools import tool

logger = logging.getLogger(__name__)


def create_export_tools(artifact_dir: str) -> list:
    """Build export tools with artifact directory injected."""

    os.makedirs(artifact_dir, exist_ok=True)

    @tool
    def export_to_xlsx(
        data: list[dict[str, Any]],
        sheet_name: str = "Sheet1",
        title: str = "export",
    ) -> dict[str, str]:
        """Export data rows to an Excel (.xlsx) file.

        Args:
            data: List of row dicts.
            sheet_name: Name for the worksheet.
            title: Base filename (without extension).

        Returns:
            Dict with file path and size.
        """
        import openpyxl
        from openpyxl.styles import Font, PatternFill

        if not data:
            return {"error": "No data to export", "path": ""}

        file_id = str(uuid.uuid4())[:8]
        path = os.path.join(artifact_dir, f"{title}_{file_id}.xlsx")

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        # Header
        headers = list(data[0].keys())
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")

        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill

        # Data rows
        for row_idx, row in enumerate(data, 2):
            for col_idx, header in enumerate(headers, 1):
                ws.cell(row=row_idx, column=col_idx, value=row.get(header))

        # Auto-width columns
        for col_idx, header in enumerate(headers, 1):
            max_len = max(
                len(str(header)),
                *(len(str(row.get(header, ""))) for row in data[:100]),
            )
            ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = min(max_len + 2, 50)

        wb.save(path)
        size = os.path.getsize(path)
        logger.info("XLSX exported: %s (%d rows, %d bytes)", path, len(data), size)
        return {"path": path, "size_bytes": size, "row_count": len(data)}

    @tool
    def export_to_csv(
        data: list[dict[str, Any]],
        title: str = "export",
    ) -> dict[str, str]:
        """Export data rows to a CSV file.

        Args:
            data: List of row dicts.
            title: Base filename.

        Returns:
            Dict with file path.
        """
        if not data:
            return {"error": "No data to export", "path": ""}

        file_id = str(uuid.uuid4())[:8]
        path = os.path.join(artifact_dir, f"{title}_{file_id}.csv")

        headers = list(data[0].keys())
        with open(path, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=headers)
            writer.writeheader()
            writer.writerows(data)

        size = os.path.getsize(path)
        return {"path": path, "size_bytes": size, "row_count": len(data)}

    @tool
    def export_to_pdf(
        title: str,
        sections: list[dict[str, str]],
        include_charts: list[str] | None = None,
    ) -> dict[str, str]:
        """Generate a PDF report with sections and optional chart images.

        Args:
            title: Report title.
            sections: List of dicts with 'heading' and 'body' keys.
            include_charts: List of PNG file paths to embed.

        Returns:
            Dict with file path.
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            Image,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )

        file_id = str(uuid.uuid4())[:8]
        path = os.path.join(artifact_dir, f"{title.replace(' ', '_')}_{file_id}.pdf")

        doc = SimpleDocTemplate(path, pagesize=A4)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "ReportTitle",
            parent=styles["Heading1"],
            fontSize=20,
            spaceAfter=20,
        )
        heading_style = styles["Heading2"]
        body_style = styles["BodyText"]

        story = []
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.3 * inch))

        for section in sections:
            if section.get("heading"):
                story.append(Paragraph(section["heading"], heading_style))
                story.append(Spacer(1, 0.1 * inch))
            if section.get("body"):
                story.append(Paragraph(section["body"], body_style))
                story.append(Spacer(1, 0.2 * inch))

        for chart_path in include_charts or []:
            if os.path.isfile(chart_path):
                img = Image(chart_path, width=6 * inch, height=3 * inch)
                story.append(img)
                story.append(Spacer(1, 0.2 * inch))

        doc.build(story)
        size = os.path.getsize(path)
        logger.info("PDF exported: %s (%d bytes)", path, size)
        return {"path": path, "size_bytes": size}

    @tool
    def export_to_docx(
        title: str,
        sections: list[dict[str, str]],
        include_charts: list[str] | None = None,
    ) -> dict[str, str]:
        """Generate a Word (.docx) document with sections and optional charts.

        Args:
            title: Document title.
            sections: List of dicts with 'heading' and 'body'.
            include_charts: PNG paths to embed.

        Returns:
            Dict with file path.
        """
        from docx import Document
        from docx.shared import Inches

        file_id = str(uuid.uuid4())[:8]
        path = os.path.join(artifact_dir, f"{title.replace(' ', '_')}_{file_id}.docx")

        doc = Document()
        doc.add_heading(title, level=0)

        for section in sections:
            if section.get("heading"):
                doc.add_heading(section["heading"], level=1)
            if section.get("body"):
                doc.add_paragraph(section["body"])

        for chart_path in include_charts or []:
            if os.path.isfile(chart_path):
                doc.add_picture(chart_path, width=Inches(6))

        doc.save(path)
        size = os.path.getsize(path)
        logger.info("DOCX exported: %s (%d bytes)", path, size)
        return {"path": path, "size_bytes": size}

    @tool
    def export_to_json(
        data: Any,
        title: str = "export",
    ) -> dict[str, str]:
        """Export data to a JSON file.

        Args:
            data: Any JSON-serializable data.
            title: Base filename.

        Returns:
            Dict with file path.
        """
        file_id = str(uuid.uuid4())[:8]
        path = os.path.join(artifact_dir, f"{title}_{file_id}.json")

        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, default=str)

        size = os.path.getsize(path)
        return {"path": path, "size_bytes": size}

    return [export_to_xlsx, export_to_csv, export_to_pdf, export_to_docx, export_to_json]
